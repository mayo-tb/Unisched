from docx import Document
import sys

def main():
    doc = Document()
    
    with open(r'C:\Users\DELL\.gemini\antigravity\brain\37c91c8a-df27-40d1-a706-cad65c688216\system_documentation.md', 'r', encoding='utf-8') as f:
        for line in f:
            line = line.strip()
            if not line:
                continue
            if line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('- **'):
                text = line[2:] # Remove '- '
                p = doc.add_paragraph(style='List Bullet')
                # Simple bolding implementation for "- **Bold**: text"
                if '**: ' in text:
                    bold_part, rest = text.split('**: ', 1)
                    bold_part = bold_part.replace('**', '')
                    r = p.add_run(bold_part + ': ')
                    r.bold = True
                    p.add_run(rest)
                else:
                    p.add_run(text.replace('**', ''))
            elif line.startswith('- '):
                p = doc.add_paragraph(line[2:], style='List Bullet')
            elif line.startswith('1.') or line.startswith('2.') or line.startswith('3.') or line.startswith('4.'):
                doc.add_paragraph(line, style='List Number')
            else:
                doc.add_paragraph(line)
                
    doc.save(r'C:\Users\DELL\unisched\Nexus_System_Documentation.docx')
    print(r'Success!')

if __name__ == '__main__':
    main()
