from docx import Document
import sys

def main():
    doc = Document()
    
    with open(r'C:\Users\DELL\.gemini\antigravity\brain\37c91c8a-df27-40d1-a706-cad65c688216\gjeta_formatted_article.md', 'r', encoding='utf-8') as f:
        for line in f:
            line = line.strip()
            if not line:
                continue
            if line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            else:
                doc.add_paragraph(line)
                
    doc.save(r'C:\Users\DELL\unisched\GJETA_Formatted_Article.docx')
    print(r"Success!")

if __name__ == '__main__':
    main()
