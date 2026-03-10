from docx import Document
from docx.shared import Inches
import sys
import os

def main():
    doc = Document()
    md_path = r'C:\Users\DELL\.gemini\antigravity\brain\37c91c8a-df27-40d1-a706-cad65c688216\updated_gjeta_article.md'
    base_dir = r'C:\Users\DELL\unisched'
    
    with open(md_path, 'r', encoding='utf-8') as f:
        for line in f:
            line = line.strip()
            if not line:
                continue
            
            if line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('- **'):
                text = line[2:]
                p = doc.add_paragraph(style='List Bullet')
                if '**: ' in text:
                    bold_part, rest = text.split('**: ', 1)
                    bold_part = bold_part.replace('**', '')
                    r = p.add_run(bold_part + ': ')
                    r.bold = True
                    p.add_run(rest)
                else:
                    p.add_run(text.replace('**', ''))
            elif line.startswith('- '):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif line.startswith('1. ') or line.startswith('2. ') or line.startswith('3. ') or line.startswith('4. ') or line.startswith('5. '):
                doc.add_paragraph(line, style='List Number')
            elif line.startswith('📸 [INSERT DIAGRAM:'):
                # Detect which diagram
                img_path = None
                if 'Architecture' in line:
                    img_path = os.path.join(base_dir, 'sys_arch.png')
                elif 'ER' in line:
                    img_path = os.path.join(base_dir, 'er_diagram.png')
                elif 'Use Case' in line:
                    img_path = os.path.join(base_dir, 'use_case.png')
                elif 'Activity' in line:
                    img_path = os.path.join(base_dir, 'activity.png')
                elif 'GA Flowchart' in line:
                    img_path = os.path.join(base_dir, 'ga_flowchart.png')
                
                if img_path and os.path.exists(img_path):
                    try:
                        doc.add_picture(img_path, width=Inches(6.0))
                        doc.paragraphs[-1].alignment = 1 # Center align
                        p = doc.add_paragraph(line) # Add the caption text below it as a normal paragraph
                        p.alignment = 1
                    except Exception as e:
                        print(f"Failed to add image {img_path}: {e}")
                        doc.add_paragraph(f"[Image placeholder missing: {img_path}]")
                else:
                    doc.add_paragraph(f"[Diagram pending generation: {img_path}]")
            else:
                doc.add_paragraph(line)
                
    out_file = os.path.join(base_dir, 'Updated_GJETA_Article_with_Diagrams.docx')
    doc.save(out_file)
    print(r'Success!')

if __name__ == '__main__':
    main()
