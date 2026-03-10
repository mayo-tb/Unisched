import sys
import subprocess
import os

try:
    import pypandoc
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "pypandoc"])
    import pypandoc

print("Downloading pandoc...")
try:
    pypandoc.download_pandoc()
except Exception as e:
    print(f"Pandoc download failed: {e}")
    # try python-docx approach if pypandoc fails
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    doc = Document()
    
    with open(r'C:\Users\DELL\.gemini\antigravity\brain\37c91c8a-df27-40d1-a706-cad65c688216\gjeta_formatted_article.md', 'r', encoding='utf-8') as f:
        for line in f:
            line = line.strip()
            if not line:
                continue
            if line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            else:
                doc.add_paragraph(line)
    doc.save(r'C:\Users\DELL\unisched\GJETA_Formatted_Article.docx')
    print("Done via python-docx.")
    sys.exit(0)

print("Converting file...")
md_path = r'C:\Users\DELL\.gemini\antigravity\brain\37c91c8a-df27-40d1-a706-cad65c688216\gjeta_formatted_article.md'
out_path = r'C:\Users\DELL\unisched\GJETA_Formatted_Article.docx'

try:
    pypandoc.convert_file(md_path, 'docx', outputfile=out_path)
    print("Done via pypandoc.")
except Exception as e:
    print(f"Conversion failed: {e}")
